from transformers import BartTokenizer, BartForConditionalGeneration
from bs4 import BeautifulSoup
import requests
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import os

# Load the BART model and tokenizer for content generation
MODEL_NAME = "facebook/bart-large-cnn"
model = BartForConditionalGeneration.from_pretrained(MODEL_NAME)
tokenizer = BartTokenizer.from_pretrained(MODEL_NAME)


def generate_chapter(topic: str, chapter_title: str) -> str:
    """
    Generates a chapter of content based on the topic and chapter title using a pre-trained model.
    Args:
        topic: Main project topic.
        chapter_title: Title of the chapter to generate.
    Returns:
        Generated text for the chapter.
    """
    prompt = f"Write a detailed, coherent chapter about '{chapter_title}' in the context of '{topic}'."
    inputs = tokenizer.encode(prompt, return_tensors="pt", max_length=512, truncation=True)
    outputs = model.generate(inputs, max_length=1024, num_beams=4, early_stopping=True)
    return tokenizer.decode(outputs[0], skip_special_tokens=True)


def search_books(topic: str, max_results: int = 3):
    """
    Searches the web for book excerpts related to the topic.
    Args:
        topic: The topic to search for.
        max_results: Number of book excerpts to retrieve.
    Returns:
        List of book excerpts or summaries.
    """
    search_url = f"https://www.google.com/search?q={topic}+book+excerpt"
    headers = {"User-Agent": "Mozilla/5.0"}
    response = requests.get(search_url, headers=headers)
    soup = BeautifulSoup(response.text, 'html.parser')

    results = []
    for i, result in enumerate(soup.find_all('div', class_='BNeawe s3v9rd AP7Wnd'), start=1):
        if i > max_results:
            break
        results.append(result.get_text())
    return results


def save_as_word_document(content: dict, references: list, filename: str):
    """
    Saves the generated content as a Word document.
    Args:
        content: Dictionary containing chapter titles and their content.
        references: List of references to include in the document.
        filename: Name of the Word document file.
    """
    doc = Document()
    doc.add_heading("Generated Project Document", level=1)

    for chapter_title, chapter_content in content.items():
        doc.add_heading(chapter_title, level=2)
        doc.add_paragraph(chapter_content)

    doc.add_heading("References", level=2)
    for ref in references:
        doc.add_paragraph(ref)

    doc.save(filename)


def save_as_pdf(content: dict, references: list, filename: str):
    """
    Saves the generated content as a PDF document.
    Args:
        content: Dictionary containing chapter titles and their content.
        references: List of references to include in the document.
        filename: Name of the PDF file.
    """
    pdf = canvas.Canvas(filename, pagesize=letter)
    width, height = letter

    pdf.setFont("Helvetica-Bold", 16)
    pdf.drawString(100, height - 50, "Generated Project Document")

    y_position = height - 100
    for chapter_title, chapter_content in content.items():
        pdf.setFont("Helvetica-Bold", 14)
        pdf.drawString(50, y_position, chapter_title)
        y_position -= 30

        pdf.setFont("Helvetica", 12)
        lines = chapter_content.split('\n')
        for line in lines:
            pdf.drawString(50, y_position, line)
            y_position -= 15
            if y_position < 50:
                pdf.showPage()
                y_position = height - 50

    pdf.setFont("Helvetica-Bold", 14)
    pdf.drawString(50, y_position - 30, "References")
    y_position -= 60

    pdf.setFont("Helvetica", 12)
    for ref in references:
        pdf.drawString(50, y_position, ref)
        y_position -= 15
        if y_position < 50:
            pdf.showPage()
            y_position = height - 50

    pdf.save()


def generate_document(topic: str, chapters: list, output_format: str = "pdf", filename: str = "generated_project"):
    """
    Main function to generate a document in the specified format.
    Args:
        topic: The main topic of the project.
        chapters: List of chapter titles.
        output_format: Output format ("pdf" or "word").
        filename: Base name of the output file (without extension).
    Returns:
        Path to the generated document file.
    """
    content = {}
    references = []

    for chapter_title in chapters:
        chapter_content = generate_chapter(topic, chapter_title)
        book_excerpts = search_books(f"{chapter_title} {topic}")

        # Integrate book excerpts into chapter content
        for excerpt in book_excerpts:
            chapter_content += f"\n\n[Excerpt] {excerpt}"
            references.append(excerpt)

        content[chapter_title] = chapter_content

    file_path = f"{filename}.{output_format}"
    if output_format.lower() == "word":
        save_as_word_document(content, references, file_path)
    elif output_format.lower() == "pdf":
        save_as_pdf(content, references, file_path)
    else:
        raise ValueError("Invalid output format. Choose 'pdf' or 'word'.")

    return file_path
